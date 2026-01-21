from docx import Document
import sys

try:
    doc = Document('/var/local/surian-ai-studio/مشروع الحكومة/مسودة ال SRS النسخة 3.docx')
    keywords = ['header', 'logo', 'color', 'navigation', 'identity', 'الترويسة', 'الشعار', 'الألوان', 'القائمة', 'nav', 'navbar', 'branding', 'اللون', 'القوائم']
    
    print("--- Relevant SRS Sections ---")
    for para in doc.paragraphs:
        text = para.text.lower()
        if any(k in text for k in keywords):
            print(para.text)
            
    print("\n--- Tables Content (often contains specs) ---")
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join([cell.text for cell in row.cells])
            if any(k in row_text.lower() for k in keywords):
                print(row_text)

except Exception as e:
    print(f"Error: {e}")
